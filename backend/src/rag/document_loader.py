from io import BytesIO
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = DocxDocument(BytesIO(file_bytes))
    return "\n".join(p.text for p in doc.paragraphs)


def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text(filename: str, file_bytes: bytes) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == "docx":
        return extract_text_from_docx(file_bytes)
    elif ext == "txt":
        return extract_text_from_txt(file_bytes)
    raise ValueError(f"Unsupported file type: .{ext}")


def chunk_text_to_documents(text: str, filename: str) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(text)
    return [
        Document(page_content=chunk, metadata={"source": filename, "title": filename})
        for chunk in chunks
    ]


def process_uploaded_file(filename: str, file_bytes: bytes) -> list[Document]:
    text = extract_text(filename, file_bytes)
    if not text.strip():
        raise ValueError("No text could be extracted from this file — it may be empty, scanned/image-only, or corrupted.")
    return chunk_text_to_documents(text, filename)